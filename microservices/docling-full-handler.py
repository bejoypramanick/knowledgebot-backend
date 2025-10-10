#!/usr/bin/env python3
"""
Docling Full - Complete Document Processing
Handles complex documents with OCR, table detection, and advanced processing
Optimized for comprehensive document analysis
"""

import json
import os
import base64
from typing import Dict, Any, List
import logging
import io
from pathlib import Path

logger = logging.getLogger(__name__)
logger.setLevel(logging.INFO)

# Initialize comprehensive libraries
try:
    import docling
    import easyocr
    import cv2
    import numpy as np
    from PIL import Image
    import pytesseract
    import pypdfium2 as pdfium
    import pdfplumber
    import fitz  # PyMuPDF
    from docx import Document
    import pandas as pd
    logger.info("Docling Full libraries initialized successfully")
except Exception as e:
    logger.error(f"Failed to initialize Docling Full libraries: {e}")
    docling = None
    easyocr = None
    cv2 = None
    np = None
    Image = None
    pytesseract = None
    pdfium = None
    pdfplumber = None
    fitz = None
    Document = None
    pd = None

def perform_ocr_on_image(image_bytes: bytes, languages: List[str] = ['en']) -> Dict[str, Any]:
    """Perform OCR on image using multiple methods"""
    results = {
        "success": False,
        "text": "",
        "detected_text": [],
        "method_used": None,
        "confidence_scores": []
    }
    
    try:
        # Convert bytes to image
        if Image and np:
            image = Image.open(io.BytesIO(image_bytes))
            image_array = np.array(image)
            
            # Method 1: EasyOCR (best for complex layouts)
            if easyocr:
                reader = easyocr.Reader(languages)
                easyocr_results = reader.readtext(image_array)
                
                text_parts = []
                confidences = []
                
                for (bbox, text, confidence) in easyocr_results:
                    text_parts.append(text)
                    confidences.append(confidence)
                
                combined_text = " ".join(text_parts)
                
                results.update({
                    "success": True,
                    "text": combined_text,
                    "detected_text": [
                        {
                            "text": text,
                            "confidence": conf,
                            "bbox": bbox
                        }
                        for (bbox, text, conf) in easyocr_results
                    ],
                    "method_used": "easyocr",
                    "confidence_scores": confidences,
                    "average_confidence": sum(confidences) / len(confidences) if confidences else 0
                })
            
            # Method 2: Tesseract (fallback)
            elif pytesseract:
                text = pytesseract.image_to_string(image)
                confidence_data = pytesseract.image_to_data(image, output_type=pytesseract.Output.DICT)
                
                confidences = [int(conf) for conf in confidence_data['conf'] if int(conf) > 0]
                
                results.update({
                    "success": True,
                    "text": text,
                    "detected_text": [
                        {
                            "text": text,
                            "confidence": conf / 100.0,
                            "bbox": None
                        }
                        for text, conf in zip(confidence_data['text'], confidence_data['conf'])
                        if text.strip() and int(conf) > 0
                    ],
                    "method_used": "tesseract",
                    "confidence_scores": [c / 100.0 for c in confidences],
                    "average_confidence": sum(confidences) / len(confidences) / 100.0 if confidences else 0
                })
            
            else:
                results["error"] = "No OCR libraries available"
        
    except Exception as e:
        logger.error(f"OCR processing failed: {e}")
        results["error"] = str(e)
    
    return results

def detect_tables_in_image(image_bytes: bytes) -> Dict[str, Any]:
    """Detect tables in image using computer vision"""
    results = {
        "success": False,
        "tables": [],
        "table_count": 0,
        "method_used": None
    }
    
    try:
        if Image and cv2 and np:
            # Convert bytes to image
            image = Image.open(io.BytesIO(image_bytes))
            image_array = np.array(image)
            
            # Convert to grayscale
            if len(image_array.shape) == 3:
                gray = cv2.cvtColor(image_array, cv2.COLOR_RGB2GRAY)
            else:
                gray = image_array
            
            # Detect horizontal and vertical lines
            horizontal_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (40, 1))
            vertical_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (1, 40))
            
            # Detect horizontal lines
            horizontal_lines = cv2.morphologyEx(gray, cv2.MORPH_OPEN, horizontal_kernel)
            horizontal_lines = cv2.dilate(horizontal_lines, horizontal_kernel, iterations=2)
            
            # Detect vertical lines
            vertical_lines = cv2.morphologyEx(gray, cv2.MORPH_OPEN, vertical_kernel)
            vertical_lines = cv2.dilate(vertical_lines, vertical_kernel, iterations=2)
            
            # Combine lines
            table_mask = cv2.addWeighted(horizontal_lines, 0.5, vertical_lines, 0.5, 0.0)
            
            # Find contours
            contours, _ = cv2.findContours(table_mask, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
            
            tables = []
            for i, contour in enumerate(contours):
                area = cv2.contourArea(contour)
                if area > 1000:  # Filter small contours
                    x, y, w, h = cv2.boundingRect(contour)
                    tables.append({
                        "table_id": i,
                        "bbox": [x, y, w, h],
                        "area": area,
                        "confidence": min(area / 10000, 1.0)  # Simple confidence based on area
                    })
            
            results.update({
                "success": True,
                "tables": tables,
                "table_count": len(tables),
                "method_used": "opencv_morphology"
            })
        
    except Exception as e:
        logger.error(f"Table detection failed: {e}")
        results["error"] = str(e)
    
    return results

def process_pdf_with_docling(document_bytes: bytes) -> Dict[str, Any]:
    """Process PDF using Docling for comprehensive analysis"""
    results = {
        "success": False,
        "text": "",
        "hierarchical_chunks": [],
        "tables": [],
        "images": [],
        "metadata": {},
        "method_used": "docling"
    }
    
    try:
        if docling:
            # Initialize Docling converter
            converter = docling.DocumentConverter()
            
            # Convert document
            doc = converter.convert(io.BytesIO(document_bytes))
            
            # Extract text
            text_content = doc.export_to_markdown()
            
            # Extract hierarchical structure
            chunks = []
            for element in doc.iterate_items():
                chunk = {
                    "content": getattr(element, 'text', '') or str(element),
                    "element_type": element.__class__.__name__,
                    "page_number": getattr(element, 'page', None),
                    "bbox": getattr(element, 'bbox', None),
                    "confidence": getattr(element, 'confidence', 1.0)
                }
                chunks.append(chunk)
            
            # Extract tables
            tables = []
            for element in doc.iterate_items():
                if element.__class__.__name__ == 'Table':
                    table_data = []
                    if hasattr(element, 'rows'):
                        for row in element.rows:
                            row_data = [cell.text if hasattr(cell, 'text') else str(cell) for cell in row]
                            table_data.append(row_data)
                    
                    tables.append({
                        "data": table_data,
                        "rows": len(table_data),
                        "cols": len(table_data[0]) if table_data else 0,
                        "page_number": getattr(element, 'page', None)
                    })
            
            # Extract images
            images = []
            for element in doc.iterate_items():
                if element.__class__.__name__ == 'Image':
                    images.append({
                        "page_number": getattr(element, 'page', None),
                        "bbox": getattr(element, 'bbox', None),
                        "caption": getattr(element, 'caption', '')
                    })
            
            results.update({
                "success": True,
                "text": text_content,
                "hierarchical_chunks": chunks,
                "tables": tables,
                "images": images,
                "metadata": {
                    "total_elements": len(chunks),
                    "total_tables": len(tables),
                    "total_images": len(images),
                    "total_text_length": len(text_content)
                }
            })
        
    except Exception as e:
        logger.error(f"Docling PDF processing failed: {e}")
        results["error"] = str(e)
    
    return results

def process_scanned_pdf(document_bytes: bytes) -> Dict[str, Any]:
    """Process scanned PDF by converting pages to images and performing OCR"""
    results = {
        "success": False,
        "text": "",
        "pages": [],
        "total_pages": 0,
        "method_used": "scanned_pdf_ocr"
    }
    
    try:
        if fitz and Image:
            # Open PDF with PyMuPDF
            pdf_document = fitz.open(stream=document_bytes, filetype="pdf")
            pages = []
            full_text = ""
            
            for page_num in range(len(pdf_document)):
                page = pdf_document[page_num]
                
                # Convert page to image
                mat = fitz.Matrix(2.0, 2.0)  # 2x zoom for better OCR
                pix = page.get_pixmap(matrix=mat)
                img_data = pix.tobytes("png")
                
                # Perform OCR on the image
                ocr_result = perform_ocr_on_image(img_data)
                
                if ocr_result["success"]:
                    page_text = ocr_result["text"]
                    pages.append({
                        "page_number": page_num + 1,
                        "text": page_text,
                        "ocr_confidence": ocr_result.get("average_confidence", 0),
                        "method_used": ocr_result.get("method_used", "unknown")
                    })
                    full_text += page_text + "\n"
            
            results.update({
                "success": True,
                "text": full_text,
                "pages": pages,
                "total_pages": len(pages),
                "method_used": "scanned_pdf_ocr"
            })
        
    except Exception as e:
        logger.error(f"Scanned PDF processing failed: {e}")
        results["error"] = str(e)
    
    return results

def process_document_full(document_bytes: bytes, filename: str, document_id: str = None, analysis: Dict[str, Any] = None) -> Dict[str, Any]:
    """
    Process document using comprehensive methods
    Handles OCR, table detection, and advanced document analysis
    """
    try:
        file_extension = Path(filename).suffix.lower()
        
        # Use analysis to determine processing strategy
        if analysis:
            is_scanned = analysis.get("is_scanned", False)
            has_tables = analysis.get("has_tables", False)
            complexity_score = analysis.get("complexity_score", 0)
        else:
            is_scanned = False
            has_tables = False
            complexity_score = 0
        
        # Route to appropriate processor
        if file_extension == '.pdf':
            if is_scanned:
                result = process_scanned_pdf(document_bytes)
            else:
                # Try Docling first, fallback to traditional methods
                result = process_pdf_with_docling(document_bytes)
                if not result["success"]:
                    # Fallback to basic PDF processing
                    import pypdfium2 as pdfium
                    if pdfium:
                        pdf = pdfium.PdfDocument(document_bytes)
                        text_content = ""
                        for i in range(len(pdf)):
                            page = pdf[i]
                            textpage = page.get_textpage()
                            text = textpage.get_text_range()
                            text_content += text + "\n"
                        
                        result = {
                            "success": True,
                            "text": text_content,
                            "method_used": "pypdfium2_fallback",
                            "metadata": {"total_pages": len(pdf)}
                        }
        
        elif file_extension in ['.jpg', '.jpeg', '.png', '.tiff', '.bmp']:
            # Process image with OCR
            result = perform_ocr_on_image(document_bytes)
            
            # Also detect tables if needed
            if has_tables:
                table_result = detect_tables_in_image(document_bytes)
                result["tables"] = table_result.get("tables", [])
                result["table_count"] = table_result.get("table_count", 0)
        
        elif file_extension in ['.docx', '.doc']:
            # Process DOCX with advanced features
            from docx import Document
            doc = Document(io.BytesIO(document_bytes))
            
            # Extract text and tables
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            tables = []
            
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)
            
            result = {
                "success": True,
                "text": "\n".join(paragraphs),
                "tables": tables,
                "method_used": "docx_advanced",
                "metadata": {
                    "total_paragraphs": len(paragraphs),
                    "total_tables": len(tables)
                }
            }
        
        else:
            # Fallback to text processing
            try:
                text = document_bytes.decode('utf-8')
                result = {
                    "success": True,
                    "text": text,
                    "method_used": "text_fallback",
                    "metadata": {"total_text_length": len(text)}
                }
            except:
                result = {
                    "success": False,
                    "error": "Unsupported file format"
                }
        
        # Add common metadata
        result.update({
            "document_id": document_id,
            "filename": filename,
            "file_extension": file_extension,
            "processing_service": "docling-full",
            "analysis": analysis
        })
        
        return result
        
    except Exception as e:
        logger.error(f"Full document processing failed: {e}")
        return {
            "success": False,
            "error": f"Full processing failed: {str(e)}",
            "document_id": document_id,
            "filename": filename,
            "processing_service": "docling-full"
        }

def lambda_handler(event: Dict[str, Any], context: Any) -> Dict[str, Any]:
    """
    Lambda handler for Docling Full processing
    """
    try:
        # Extract document data
        document_bytes = base64.b64decode(event.get('document_bytes', ''))
        filename = event.get('filename', 'unknown')
        document_id = event.get('document_id')
        analysis = event.get('analysis', {})
        
        if not document_bytes:
            return {
                "statusCode": 400,
                "body": json.dumps({
                    "success": False,
                    "error": "No document bytes provided"
                })
            }
        
        # Process document
        import time
        start_time = time.time()
        result = process_document_full(document_bytes, filename, document_id, analysis)
        result["processing_time"] = time.time() - start_time
        
        return {
            "statusCode": 200,
            "body": json.dumps(result)
        }
        
    except Exception as e:
        logger.error(f"Lambda handler error: {e}")
        return {
            "statusCode": 500,
            "body": json.dumps({
                "success": False,
                "error": f"Full processing failed: {str(e)}"
            })
        }

