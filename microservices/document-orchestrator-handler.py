#!/usr/bin/env python3
"""
Document Orchestrator - Intelligent Document Processing
Routes documents to appropriate services based on analysis
Prevents duplicate processing and optimizes resource usage
"""

import json
import os
import base64
import io
from typing import Dict, Any, List
import logging
from pathlib import Path

logger = logging.getLogger(__name__)
logger.setLevel(logging.INFO)

# Initialize analysis libraries
try:
    import magic
    from PIL import Image
    import cv2
    import numpy as np
    logger.info("Document analysis libraries initialized successfully")
except Exception as e:
    logger.error(f"Failed to initialize analysis libraries: {e}")
    magic = None
    Image = None
    cv2 = None
    np = None

async def call_microservice(url: str, payload: Dict[str, Any]) -> Dict[str, Any]:
    """Call a microservice with error handling"""
    import aiohttp
    async with aiohttp.ClientSession() as client:
        try:
            response = await client.post(url, json=payload)
            return await response.json()
        except Exception as e:
            logger.error(f"Failed to call microservice {url}: {e}")
            return {"success": False, "error": str(e)}

def analyze_document_complexity(document_bytes: bytes, filename: str) -> Dict[str, Any]:
    """
    Analyze document to determine processing requirements
    Returns analysis with recommended processing strategy
    """
    analysis = {
        "filename": filename,
        "file_type": "unknown",
        "is_scanned": False,
        "has_tables": False,
        "has_images": False,
        "complexity_score": 0,
        "recommended_service": "docling-core",
        "processing_steps": []
    }
    
    try:
        # Determine file type
        if magic:
            mime_type = magic.from_buffer(document_bytes, mime=True)
            analysis["file_type"] = mime_type
        else:
            # Fallback to file extension
            ext = Path(filename).suffix.lower()
            if ext == '.pdf':
                analysis["file_type"] = 'application/pdf'
            elif ext in ['.jpg', '.jpeg', '.png', '.tiff', '.bmp']:
                analysis["file_type"] = 'image'
            elif ext in ['.docx', '.doc']:
                analysis["file_type"] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        
        # Analyze based on file type
        if analysis["file_type"] == 'application/pdf':
            analysis.update(_analyze_pdf_complexity(document_bytes))
        elif analysis["file_type"].startswith('image/'):
            analysis.update(_analyze_image_complexity(document_bytes))
        elif analysis["file_type"].endswith('wordprocessingml.document'):
            analysis.update(_analyze_docx_complexity(document_bytes))
        else:
            # Default analysis for unknown types
            analysis["complexity_score"] = 1
            analysis["recommended_service"] = "docling-core"
            analysis["processing_steps"] = ["basic_text_extraction"]
        
        # Determine final recommendation
        if analysis["complexity_score"] >= 7:
            analysis["recommended_service"] = "docling-full"
        elif analysis["complexity_score"] >= 4:
            analysis["recommended_service"] = "docling-core"
        else:
            analysis["recommended_service"] = "docling-core"
            
        logger.info(f"Document analysis complete: {analysis}")
        return analysis
        
    except Exception as e:
        logger.error(f"Failed to analyze document: {e}")
        return {
            "filename": filename,
            "file_type": "unknown",
            "complexity_score": 1,
            "recommended_service": "docling-core",
            "processing_steps": ["basic_text_extraction"],
            "error": str(e)
        }

def _analyze_pdf_complexity(document_bytes: bytes) -> Dict[str, Any]:
    """Analyze PDF complexity"""
    analysis = {
        "complexity_score": 0,
        "processing_steps": []
    }
    
    try:
        # Try to extract text to check if it's scanned
        import pypdfium2 as pdfium
        pdf = pdfium.PdfDocument(document_bytes)
        
        # Check first few pages for text content
        text_content = ""
        for i in range(min(3, len(pdf))):
            page = pdf[i]
            textpage = page.get_textpage()
            text = textpage.get_text_range()
            text_content += text
        
        # Analyze text content
        if len(text_content.strip()) < 100:
            analysis["is_scanned"] = True
            analysis["complexity_score"] += 3
            analysis["processing_steps"].append("ocr_required")
        else:
            analysis["complexity_score"] += 1
            analysis["processing_steps"].append("text_extraction")
        
        # Check for images (simplified check)
        if len(document_bytes) > 5 * 1024 * 1024:  # > 5MB
            analysis["has_images"] = True
            analysis["complexity_score"] += 2
            analysis["processing_steps"].append("image_processing")
        
        # Check for potential tables (heuristic)
        if any(keyword in text_content.lower() for keyword in ['table', 'figure', 'chart', 'graph']):
            analysis["has_tables"] = True
            analysis["complexity_score"] += 2
            analysis["processing_steps"].append("table_detection")
        
    except Exception as e:
        logger.warning(f"PDF analysis failed: {e}")
        analysis["complexity_score"] = 2
        analysis["processing_steps"] = ["basic_pdf_processing"]
    
    return analysis

def _analyze_image_complexity(document_bytes: bytes) -> Dict[str, Any]:
    """Analyze image complexity"""
    analysis = {
        "complexity_score": 0,
        "processing_steps": []
    }
    
    try:
        if Image and cv2 and np:
            # Convert to image
            image = Image.open(io.BytesIO(document_bytes))
            image_array = np.array(image)
            
            # Basic complexity analysis
            analysis["complexity_score"] = 3  # Images always need OCR
            analysis["processing_steps"].append("ocr_required")
            
            # Check image size
            height, width = image_array.shape[:2]
            if height > 2000 or width > 2000:
                analysis["complexity_score"] += 1
                analysis["processing_steps"].append("high_resolution_processing")
            
            # Check for potential tables (simplified)
            gray = cv2.cvtColor(image_array, cv2.COLOR_RGB2GRAY) if len(image_array.shape) == 3 else image_array
            edges = cv2.Canny(gray, 50, 150)
            lines = cv2.HoughLinesP(edges, 1, np.pi/180, threshold=100, minLineLength=100, maxLineGap=10)
            
            if lines is not None and len(lines) > 20:
                analysis["has_tables"] = True
                analysis["complexity_score"] += 2
                analysis["processing_steps"].append("table_detection")
        
    except Exception as e:
        logger.warning(f"Image analysis failed: {e}")
        analysis["complexity_score"] = 3
        analysis["processing_steps"] = ["basic_ocr"]
    
    return analysis

def _analyze_docx_complexity(document_bytes: bytes) -> Dict[str, Any]:
    """Analyze DOCX complexity"""
    analysis = {
        "complexity_score": 0,
        "processing_steps": []
    }
    
    try:
        from docx import Document
        doc = Document(io.BytesIO(document_bytes))
        
        # Count elements
        paragraph_count = len(doc.paragraphs)
        table_count = len(doc.tables)
        
        analysis["complexity_score"] = 1  # Base complexity
        analysis["processing_steps"].append("text_extraction")
        
        if table_count > 0:
            analysis["has_tables"] = True
            analysis["complexity_score"] += 2
            analysis["processing_steps"].append("table_extraction")
        
        if paragraph_count > 50:
            analysis["complexity_score"] += 1
            analysis["processing_steps"].append("large_document_processing")
        
    except Exception as e:
        logger.warning(f"DOCX analysis failed: {e}")
        analysis["complexity_score"] = 2
        analysis["processing_steps"] = ["basic_docx_processing"]
    
    return analysis

async def process_document_intelligently(document_bytes: bytes, filename: str, document_id: str = None) -> Dict[str, Any]:
    """
    Intelligently process document using appropriate services
    Prevents duplicate processing and optimizes resource usage
    """
    try:
        # Step 1: Analyze document complexity
        analysis = analyze_document_complexity(document_bytes, filename)
        
        logger.info(f"Document analysis: {analysis}")
        
        # Step 2: Route to appropriate service based on analysis
        recommended_service = analysis["recommended_service"]
        
        if recommended_service == "docling-full":
            # Use docling-full for complex documents
            result = await call_microservice(
                f"{os.environ.get('DOCLING_FULL_SERVICE_URL', 'http://localhost:8080')}/docling-full",
                {
                    "document_bytes": base64.b64encode(document_bytes).decode(),
                    "filename": filename,
                    "document_id": document_id,
                    "analysis": analysis
                }
            )
        else:
            # Use docling-core for simple documents
            result = await call_microservice(
                f"{os.environ.get('DOCLING_CORE_SERVICE_URL', 'http://localhost:8080')}/docling-core",
                {
                    "document_bytes": base64.b64encode(document_bytes).decode(),
                    "filename": filename,
                    "document_id": document_id,
                    "analysis": analysis
                }
            )
        
        # Step 3: Add analysis metadata to result
        if result.get("success"):
            result["analysis"] = analysis
            result["processing_strategy"] = recommended_service
            result["processing_steps"] = analysis["processing_steps"]
        
        return result
        
    except Exception as e:
        logger.error(f"Intelligent document processing failed: {e}")
        return {
            "success": False,
            "error": f"Document processing failed: {str(e)}",
            "document_id": document_id,
            "analysis": analysis if 'analysis' in locals() else None
        }

def lambda_handler(event: Dict[str, Any], context: Any) -> Dict[str, Any]:
    """
    Lambda handler for intelligent document processing
    """
    try:
        # Extract document data
        document_bytes = base64.b64decode(event.get('document_bytes', ''))
        filename = event.get('filename', 'unknown')
        document_id = event.get('document_id')
        
        if not document_bytes:
            return {
                "statusCode": 400,
                "body": json.dumps({
                    "success": False,
                    "error": "No document bytes provided"
                })
            }
        
        # Process document intelligently
        import asyncio
        result = asyncio.run(process_document_intelligently(document_bytes, filename, document_id))
        
        return {
            "statusCode": 200,
            "body": json.dumps(result)
        }
        
    except Exception as e:
        logger.error(f"Lambda handler error: {e}")
        return {
            "statusCode": 500,
            "body": json.dumps({
                "success": False,
                "error": f"Processing failed: {str(e)}"
            })
        }

